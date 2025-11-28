import streamlit as st
from langchain_core.documents import Document
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_openai import OpenAI
import pandas as pd
import PyPDF2
import docx
import os
import re
from typing import List, Dict
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

class ResumeScreeningAgent:
    def __init__(self, openai_api_key: str = None):
        self.openai_api_key = openai_api_key
        if openai_api_key:
            os.environ["OPENAI_API_KEY"] = openai_api_key
        
        # Use better FREE embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-mpnet-base-v2"  # Better model
        )
        self.vector_store = None
        self.resumes_data = []
        
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def process_resumes(self, uploaded_files) -> List[Document]:
        """Process uploaded resume files and extract text"""
        documents = []
        self.resumes_data = []
        
        for uploaded_file in uploaded_files:
            try:
                text = ""
                if uploaded_file.type == "application/pdf":
                    text = self.extract_text_from_pdf(uploaded_file)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    text = self.extract_text_from_docx(uploaded_file)
                else:
                    st.warning(f"Unsupported file type: {uploaded_file.type}")
                    continue
                
                if not text.strip():
                    st.warning(f"No text extracted from {uploaded_file.name}")
                    continue
                
                # Clean and normalize text
                text = self.clean_text(text)
                
                # Create document for vector store
                doc = Document(
                    page_content=text,
                    metadata={
                        "filename": uploaded_file.name,
                        "type": uploaded_file.type,
                        "size": uploaded_file.size
                    }
                )
                documents.append(doc)
                self.resumes_data.append({
                    "filename": uploaded_file.name,
                    "text": text,
                    "metadata": doc.metadata
                })
                
                st.success(f"Processed: {uploaded_file.name}")
                
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        
        return documents
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep relevant ones
        text = re.sub(r'[^\w\s\.\,\-\+\&]', ' ', text)
        return text.strip()
    
    def extract_key_skills(self, text: str) -> List[str]:
        """Extract key skills from text"""
        skills_keywords = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust'],
            'web': ['html', 'css', 'react', 'angular', 'vue', 'django', 'flask', 'node', 'express'],
            'data': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch'],
            'ml_ai': ['machine learning', 'deep learning', 'ai', 'tensorflow', 'pytorch', 'scikit-learn', 'nlp', 'computer vision'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform'],
            'tools': ['git', 'jenkins', 'jira', 'confluence', 'slack'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem solving', 'project management']
        }
        
        found_skills = []
        text_lower = text.lower()
        
        for category, skills in skills_keywords.items():
            for skill in skills:
                if skill in text_lower:
                    found_skills.append(skill)
        
        return found_skills
    
    def calculate_semantic_similarity(self, job_description: str, resume_text: str) -> float:
        """Calculate semantic similarity using TF-IDF and cosine similarity"""
        try:
            # Create TF-IDF vectors
            vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
            tfidf_matrix = vectorizer.fit_transform([job_description, resume_text])
            
            # Calculate cosine similarity
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return similarity * 100  # Convert to percentage
        except:
            return 0
    
    def advanced_scoring(self, job_description: str, resume_text: str) -> Dict:
        """Advanced scoring with multiple factors"""
        job_lower = job_description.lower()
        resume_lower = resume_text.lower()
        
        # Factor 1: Skills match
        job_skills = self.extract_key_skills(job_description)
        resume_skills = self.extract_key_skills(resume_text)
        matched_skills = set(job_skills) & set(resume_skills)
        skills_score = len(matched_skills) / max(1, len(job_skills)) * 40  # 40% weight
        
        # Factor 2: Semantic similarity
        semantic_score = self.calculate_semantic_similarity(job_description, resume_text) * 0.4  # 40% weight
        
        # Factor 3: Experience level (simple heuristic)
        exp_keywords = ['years', 'experience', 'experienced', 'senior', 'lead', 'manager']
        exp_matches = sum(1 for keyword in exp_keywords if keyword in resume_lower and keyword in job_lower)
        exp_score = (exp_matches / max(1, len(exp_keywords))) * 20  # 20% weight
        
        total_score = min(100, skills_score + semantic_score + exp_score)
        
        # Generate recommendations
        if total_score >= 80:
            recommendation = "Strong Yes"
        elif total_score >= 60:
            recommendation = "Yes"
        elif total_score >= 40:
            recommendation = "Maybe"
        else:
            recommendation = "No"
        
        strengths = [f"• Matches {len(matched_skills)} key skills" if matched_skills else "• Basic qualifications met"]
        if semantic_score > 30:
            strengths.append("• Good semantic match with job description")
        if exp_score > 10:
            strengths.append("• Relevant experience level")
        
        missing = []
        if len(matched_skills) < len(job_skills) / 2:
            missing.append("• Missing several key skills")
        if semantic_score < 20:
            missing.append("• Low semantic similarity with job requirements")
        
        return {
            'score': int(total_score),
            'strengths': '\n'.join(strengths),
            'missing': '\n'.join(missing) if missing else "• All key areas covered",
            'recommendation': recommendation,
            'skills_match': f"• {len(matched_skills)}/{len(job_skills)} skills matched: {', '.join(list(matched_skills)[:5])}"
        }
    
    def create_vector_store(self, documents: List[Document]):
        """Create vector store from resume documents"""
        if not documents:
            return None
        
        # Split documents into chunks
        text_splitter = CharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        texts = text_splitter.split_documents(documents)
        
        # Create vector store with FREE embeddings
        self.vector_store = FAISS.from_documents(texts, self.embeddings)
        return self.vector_store
    
    def screen_resume(self, job_description: str, resume_text: str) -> Dict:
        """Screen individual resume against job description"""
        
        if self.openai_api_key:
            try:
                return self._ai_enhanced_screening(job_description, resume_text)
            except Exception as e:
                st.warning(f"OpenAI API error, using advanced scoring: {str(e)}")
        
        # Use advanced scoring as fallback
        return self.advanced_scoring(job_description, resume_text)
    
    def _ai_enhanced_screening(self, job_description: str, resume_text: str) -> Dict:
        """AI-enhanced screening when API key is available"""
        prompt = f"""
        You are an expert resume screener. Analyze this resume against the job description and provide accurate scoring.
        
        JOB DESCRIPTION:
        {job_description}
        
        RESUME CONTENT:
        {resume_text[:4000]}
        
        Provide a comprehensive analysis with:
        1. Match Score (0-100) - based on skills, experience, and qualifications match
        2. Key Strengths (3-5 specific points)
        3. Missing Qualifications (3-5 specific points) 
        4. Overall Recommendation (Strong Yes, Yes, Maybe, No)
        5. Key Skills Match (specific skills that match)
        
        Be strict and accurate in your assessment. Consider:
        - Required skills match
        - Experience level
        - Education qualifications
        - Relevant projects/achievements
        
        Format exactly as:
        Match Score: [number]
        Key Strengths: [point1|point2|point3]
        Missing Qualifications: [point1|point2|point3]
        Overall Recommendation: [Strong Yes|Yes|Maybe|No]
        Key Skills Match: [skill1|skill2|skill3]
        """
        
        llm = OpenAI(temperature=0.1, max_tokens=1500)
        response = llm.invoke(prompt)
        
        return self._parse_ai_response(response)
    
    def _parse_ai_response(self, response: str) -> Dict:
        """Parse AI response into structured data"""
        result = {}
        lines = response.split('\n')
        for line in lines:
            if 'Match Score:' in line:
                try:
                    result['score'] = int(line.split(':')[1].strip().split()[0])
                except:
                    result['score'] = 50
            elif 'Key Strengths:' in line:
                result['strengths'] = line.split(':')[1].strip().replace('|', '\n• ')
            elif 'Missing Qualifications:' in line:
                result['missing'] = line.split(':')[1].strip().replace('|', '\n• ')
            elif 'Overall Recommendation:' in line:
                result['recommendation'] = line.split(':')[1].strip()
            elif 'Key Skills Match:' in line:
                result['skills_match'] = line.split(':')[1].strip().replace('|', '\n• ')
        
        # Ensure all keys exist
        result.setdefault('score', 50)
        result.setdefault('strengths', '• Analysis in progress')
        result.setdefault('missing', '• Analysis in progress')
        result.setdefault('recommendation', 'Maybe')
        result.setdefault('skills_match', '• Skills analysis in progress')
        
        return result
    
    def rank_resumes(self, job_description: str) -> List[Dict]:
        """Rank all resumes based on job description match"""
        if not self.resumes_data:
            return []
        
        ranked_resumes = []
        progress_bar = st.progress(0)
        
        for i, resume in enumerate(self.resumes_data):
            progress_bar.progress((i + 1) / len(self.resumes_data))
            with st.spinner(f"Analyzing {resume['filename']}..."):
                result = self.screen_resume(job_description, resume['text'])
                if result:
                    ranked_resumes.append({
                        'filename': resume['filename'],
                        'score': result.get('score', 0),
                        'recommendation': result.get('recommendation', 'Unknown'),
                        'strengths': result.get('strengths', ''),
                        'missing': result.get('missing', ''),
                        'skills_match': result.get('skills_match', '')
                    })
        
        # Sort by score descending
        ranked_resumes.sort(key=lambda x: x['score'], reverse=True)
        return ranked_resumes

def main():
    st.set_page_config(
        page_title="Advanced Resume Screening Agent",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("🤖 Advanced Resume Screening Agent")
    st.markdown("""
    **More accurate resume screening with:**
    - Semantic similarity analysis
    - Advanced keyword matching  
    - Skills extraction
    - Experience level assessment
    - Optional AI-enhanced analysis
    """)
    
    # Sidebar for API key and settings
    with st.sidebar:
        st.header("Configuration")
        openai_api_key = st.text_input("OpenAI API Key (Optional)", type="password", 
                                      help="Add for AI-powered analysis, leave empty for advanced scoring")
        
        st.markdown("---")
        st.info("""
        **Accuracy Features:**
        - ✅ Semantic similarity scoring
        - ✅ Skills extraction & matching
        - ✅ Experience level assessment
        - 🔄 AI-enhanced analysis (with API key)
        - 📊 Multi-factor scoring
        """)
    
    # Initialize the screening agent
    screening_agent = ResumeScreeningAgent(openai_api_key)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📝 Job Description")
        job_description = st.text_area(
            "Enter detailed job description:",
            height=250,
            placeholder="Include:\n• Required skills and technologies\n• Experience level\n• Education requirements\n• Key responsibilities\n• Preferred qualifications"
        )
    
    with col2:
        st.subheader("📄 Upload Resumes")
        uploaded_files = st.file_uploader(
            "Choose resume files",
            type=['pdf', 'docx'],
            accept_multiple_files=True,
            help="Upload multiple PDF or DOCX files"
        )
        
        if uploaded_files:
            st.success(f"📁 {len(uploaded_files)} file(s) uploaded")
            
        # Analysis type info
        if openai_api_key:
            st.info("🔮 **AI-Powered Analysis**: Using OpenAI for detailed assessment")
        else:
            st.info("📊 **Advanced Scoring**: Using semantic analysis + skills matching")
    
    # Process resumes when files are uploaded
    if uploaded_files and job_description:
        if st.button("🚀 Start Advanced Screening", type="primary"):
            with st.spinner("Processing resumes with advanced analysis..."):
                # Process resumes
                documents = screening_agent.process_resumes(uploaded_files)
                
                if documents:
                    st.success(f"✅ Successfully processed {len(documents)} resume(s)")
                    
                    # Create vector store
                    with st.spinner("Building semantic search index..."):
                        screening_agent.create_vector_store(documents)
                    
                    # Rank resumes
                    ranked_resumes = screening_agent.rank_resumes(job_description)
                    
                    # Display results
                    st.subheader("🎯 Advanced Screening Results")
                    
                    # Create results dataframe
                    results_data = []
                    for resume in ranked_resumes:
                        results_data.append({
                            'Rank': len(results_data) + 1,
                            'Filename': resume['filename'],
                            'Match Score': resume['score'],
                            'Recommendation': resume['recommendation']
                        })
                    
                    if results_data:
                        df = pd.DataFrame(results_data)
                        st.dataframe(df, use_container_width=True)
                        
                        # Summary metrics
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("Total Candidates", len(ranked_resumes))
                        with col2:
                            strong_matches = len([r for r in ranked_resumes if r['score'] >= 80])
                            st.metric("Strong Matches", strong_matches)
                        with col3:
                            avg_score = sum([r['score'] for r in ranked_resumes]) / len(ranked_resumes)
                            st.metric("Average Score", f"{avg_score:.1f}/100")
                        with col4:
                            qualified = len([r for r in ranked_resumes if r['score'] >= 60])
                            st.metric("Qualified", qualified)
                        
                        # Detailed view
                        st.subheader("📊 Detailed Analysis")
                        
                        for i, resume in enumerate(ranked_resumes):
                            with st.expander(f"{i+1}. {resume['filename']} - Score: {resume['score']}/100 - {resume['recommendation']}"):
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.write("**✅ Key Strengths:**")
                                    st.write(resume['strengths'])
                                    
                                    st.write("**🔍 Skills Match:**")
                                    st.write(resume['skills_match'])
                                
                                with col2:
                                    st.write("**❌ Missing Qualifications:**")
                                    st.write(resume['missing'])
                                    
                                    st.write("**🎯 Recommendation:**")
                                    rec_color = {
                                        'Strong Yes': 'green',
                                        'Yes': 'lightgreen',
                                        'Maybe': 'orange',
                                        'No': 'red'
                                    }.get(resume['recommendation'], 'gray')
                                    st.markdown(f"<span style='color: {rec_color}; font-weight: bold; font-size: 16px;'>{resume['recommendation']}</span>", 
                                              unsafe_allow_html=True)
                        
                        # Export option
                        st.subheader("📤 Export Results")
                        if st.button("💾 Download Detailed Results as CSV"):
                            df_export = pd.DataFrame(ranked_resumes)
                            csv = df_export.to_csv(index=False)
                            st.download_button(
                                label="Click to Download CSV",
                                data=csv,
                                file_name="advanced_screening_results.csv",
                                mime="text/csv"
                            )
                
                else:
                    st.error("❌ No text could be extracted from the uploaded files.")

if __name__ == "__main__":
    main()